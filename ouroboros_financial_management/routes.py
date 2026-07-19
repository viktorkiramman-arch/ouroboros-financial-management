from __future__ import annotations

import io
from datetime import date
from decimal import Decimal

from docx import Document
from flask import Blueprint, abort, flash, g, jsonify, redirect, render_template, request, send_file, url_for
from openpyxl import Workbook
from sqlalchemy.exc import IntegrityError

from .constants import (
    ANIMATION_LEVELS,
    APPEARANCE_OPTIONS,
    CATEGORY_CHARTS,
    COLOR_THEMES,
    CURRENCY_OPTIONS,
    INCOME_EXPENSE_CHARTS,
    MEMBER_ROLES,
    WEEK_STARTS,
    WORKSPACE_TYPES,
)
from .dates import next_month, previous_month
from .extensions import db
from .models import AnalystRun, Budget, Category, Rule, Transaction, UserSetting, Workspace, WorkspaceMember
from .money import ValidationError, clean_text, normalize_category, parse_money_to_cents, spreadsheet_safe
from .security import current_user_id, login_required, rate_limit
from .services import (
    budget_status_rows,
    calendar_data,
    dashboard_summary,
    ensure_active_workspace,
    ensure_user_settings,
    generate_analyst_report,
    get_exchange_rate,
    import_csv_text,
    load_csv_preview_text,
    ouroboros_advisor_reply,
    parse_analyst_result,
    preview_csv,
    save_csv_preview_text,
    set_active_workspace,
    upsert_category,
    validate_member_form,
    validate_transaction_form,
    validate_workspace_form,
)

bp = Blueprint("main", __name__)


@bp.before_request
def load_settings_and_workspace():
    if g.get("user"):
        g.settings = ensure_user_settings(g.user)
        g.workspace = ensure_active_workspace(g.user)
        # Persist any auto-created workspace/settings from first login without waiting for a write route.
        db.session.commit()
    else:
        g.settings = None
        g.workspace = None


def _safe_redirect(endpoint: str, **values):
    return redirect(url_for(endpoint, **values))


def _get_owned(model, object_id: int):
    obj = db.session.get(model, object_id)
    if obj is None or obj.user_id != current_user_id():
        abort(404)
    if hasattr(obj, "workspace_id") and obj.workspace_id not in {None, g.workspace.id}:
        abort(404)
    return obj


def _member_rows():
    return (
        WorkspaceMember.query.filter_by(user_id=current_user_id(), workspace_id=g.workspace.id)
        .order_by(WorkspaceMember.role.asc(), WorkspaceMember.name.asc())
        .all()
    )


def _transaction_json(rows: list[Transaction]) -> list[dict]:
    return [
        {
            "id": t.id,
            "date": t.date.isoformat(),
            "description": t.description,
            "category": t.category,
            "member": t.member.name if t.member else "Unassigned",
            "type": "Income" if t.is_income else "Expense",
            "is_income": bool(t.is_income),
            "amount_cents": int(t.amount_cents),
            "edit_url": url_for("main.edit_transaction", tx_id=t.id),
            "delete_url": url_for("main.delete_transaction", tx_id=t.id),
        }
        for t in rows
    ]


@bp.route("/")
@login_required
def dashboard():
    summary = dashboard_summary(current_user_id(), g.workspace.id, settings=g.settings, workspace=g.workspace)
    return render_template("dashboard.html", summary=summary, page_title="Dashboard")


@bp.route("/settings", methods=["GET", "POST"])
@login_required
@rate_limit(limit=30, window_seconds=60, label="settings")
def settings():
    setting: UserSetting = g.settings
    if request.method == "POST":
        appearance = request.form.get("appearance", "system")
        color_theme = request.form.get("color_theme", "ocean")
        income_chart = request.form.get("income_expense_chart", "line")
        category_chart = request.form.get("category_chart", "donut")
        animation_level = request.form.get("animation_level", "standard")
        week_start = request.form.get("calendar_week_start", "sunday")
        base_currency = request.form.get("base_currency_code", "USD")
        display_currency = request.form.get("display_currency_code", base_currency)
        if appearance not in APPEARANCE_OPTIONS:
            abort(400)
        if color_theme not in COLOR_THEMES:
            abort(400)
        if income_chart not in INCOME_EXPENSE_CHARTS:
            abort(400)
        if category_chart not in CATEGORY_CHARTS:
            abort(400)
        if animation_level not in ANIMATION_LEVELS:
            abort(400)
        if week_start not in WEEK_STARTS:
            abort(400)
        if base_currency not in CURRENCY_OPTIONS or display_currency not in CURRENCY_OPTIONS:
            abort(400)
        setting.appearance = appearance
        setting.color_theme = color_theme
        setting.income_expense_chart = income_chart
        setting.category_chart = category_chart
        setting.animation_level = animation_level
        setting.calendar_week_start = week_start
        setting.base_currency_code = base_currency
        setting.display_currency_code = display_currency
        setting.currency_code = base_currency
        setting.currency_symbol = CURRENCY_OPTIONS[base_currency]
        db.session.commit()
        flash("Settings saved.", "success")
        return _safe_redirect("main.settings")
    return render_template(
        "settings.html",
        page_title="Settings",
        appearance_options=sorted(APPEARANCE_OPTIONS),
        color_themes=sorted(COLOR_THEMES),
        income_charts=sorted(INCOME_EXPENSE_CHARTS),
        category_charts=sorted(CATEGORY_CHARTS),
        animation_levels=sorted(ANIMATION_LEVELS),
        week_starts=sorted(WEEK_STARTS),
        currency_options=CURRENCY_OPTIONS,
    )


@bp.route("/workspaces", methods=["GET", "POST"])
@login_required
@rate_limit(limit=40, window_seconds=60, label="workspaces")
def workspaces():
    user_id = current_user_id()
    if request.method == "POST":
        try:
            workspace = validate_workspace_form(request.form, user_id=user_id)
            db.session.add(workspace)
            db.session.flush()
            db.session.add(
                WorkspaceMember(
                    user_id=user_id,
                    workspace_id=workspace.id,
                    name=g.user.username,
                    role="breadwinner" if workspace.workspace_type in {"personal", "family"} else "member",
                    relationship="Self",
                )
            )
            set_active_workspace(user_id, workspace.id)
            db.session.commit()
            flash("Workspace created and selected.", "success")
        except (ValidationError, IntegrityError) as exc:
            db.session.rollback()
            flash(str(exc), "error")
        return _safe_redirect("main.workspaces")
    rows = Workspace.query.filter_by(user_id=user_id).order_by(Workspace.created_at.asc()).all()
    return render_template("workspaces.html", rows=rows, workspace_types=sorted(WORKSPACE_TYPES), page_title="Workspaces")


@bp.route("/workspaces/<int:workspace_id>/switch", methods=["POST"])
@login_required
@rate_limit(limit=60, window_seconds=60, label="workspace_switch")
def switch_workspace(workspace_id: int):
    try:
        set_active_workspace(current_user_id(), workspace_id)
        db.session.commit()
        flash("Workspace switched.", "success")
    except ValidationError as exc:
        db.session.rollback()
        flash(str(exc), "error")
    return _safe_redirect("main.dashboard")


@bp.route("/workspaces/<int:workspace_id>/delete", methods=["POST"])
@login_required
@rate_limit(limit=20, window_seconds=60, label="workspace_delete")
def delete_workspace(workspace_id: int):
    user_id = current_user_id()
    workspace = Workspace.query.filter_by(id=workspace_id, user_id=user_id).first_or_404()
    count = Workspace.query.filter_by(user_id=user_id).count()
    if count <= 1:
        flash("You must keep at least one workspace.", "error")
        return _safe_redirect("main.workspaces")
    AnalystRun.query.filter_by(user_id=user_id, workspace_id=workspace.id).delete()
    from .models import Insight

    Insight.query.filter_by(user_id=user_id, workspace_id=workspace.id).delete()
    Transaction.query.filter_by(user_id=user_id, workspace_id=workspace.id).delete()
    Budget.query.filter_by(user_id=user_id, workspace_id=workspace.id).delete()
    WorkspaceMember.query.filter_by(user_id=user_id, workspace_id=workspace.id).delete()
    db.session.delete(workspace)
    db.session.flush()
    first = Workspace.query.filter_by(user_id=user_id).order_by(Workspace.created_at.asc()).first()
    if first:
        g.settings.active_workspace_id = first.id
    db.session.commit()
    flash("Workspace deleted.", "success")
    return _safe_redirect("main.workspaces")


@bp.route("/members", methods=["GET", "POST"])
@login_required
@rate_limit(limit=60, window_seconds=60, label="members")
def members():
    user_id = current_user_id()
    if request.method == "POST":
        try:
            member = validate_member_form(request.form, user_id=user_id, workspace_id=g.workspace.id)
            db.session.add(member)
            db.session.commit()
            flash("Member saved.", "success")
        except ValidationError as exc:
            db.session.rollback()
            flash(str(exc), "error")
        return _safe_redirect("main.members")
    rows = _member_rows()
    return render_template("members.html", rows=rows, member_roles=sorted(MEMBER_ROLES), page_title="Members")


@bp.route("/members/<int:member_id>/delete", methods=["POST"])
@login_required
@rate_limit(limit=50, window_seconds=60, label="member_delete")
def delete_member(member_id: int):
    member: WorkspaceMember = _get_owned(WorkspaceMember, member_id)
    used = Transaction.query.filter_by(user_id=current_user_id(), workspace_id=g.workspace.id, member_id=member.id).first()
    if used:
        flash("This member has linked transactions. Remove or reassign those transactions first.", "error")
        return _safe_redirect("main.members")
    db.session.delete(member)
    db.session.commit()
    flash("Member deleted.", "success")
    return _safe_redirect("main.members")


@bp.route("/transactions", methods=["GET", "POST"])
@login_required
@rate_limit(limit=80, window_seconds=60, label="transactions")
def transactions():
    user_id = current_user_id()
    if request.method == "POST":
        try:
            tx = validate_transaction_form(request.form, user_id=user_id, workspace_id=g.workspace.id)
            db.session.add(tx)
            upsert_category(user_id, tx.category)
            db.session.commit()
            flash("Transaction saved.", "success")
        except ValidationError as exc:
            db.session.rollback()
            flash(str(exc), "error")
        return _safe_redirect("main.transactions")

    q = clean_text(request.args.get("q"), max_length=60, field_name="Search") if request.args.get("q") else ""
    category = clean_text(request.args.get("category"), max_length=40, field_name="Category filter") if request.args.get("category") else ""
    query = Transaction.query.filter_by(user_id=user_id, workspace_id=g.workspace.id)
    if q:
        query = query.filter(Transaction.description.ilike(f"%{q}%"))
    if category:
        query = query.filter(Transaction.category == category)
    rows = query.order_by(Transaction.date.desc(), Transaction.id.desc()).limit(5000).all()
    categories = Category.query.filter_by(user_id=user_id).order_by(Category.name.asc()).all()
    return render_template(
        "transactions.html",
        rows=rows,
        rows_json=_transaction_json(rows),
        categories=categories,
        members=_member_rows(),
        q=q,
        category=category,
        page_title="Transactions",
    )


@bp.route("/transactions/<int:tx_id>/edit", methods=["GET", "POST"])
@login_required
@rate_limit(limit=60, window_seconds=60, label="transaction_edit")
def edit_transaction(tx_id: int):
    tx: Transaction = _get_owned(Transaction, tx_id)
    if request.method == "POST":
        try:
            updated = validate_transaction_form(request.form, user_id=current_user_id(), workspace_id=g.workspace.id)
            tx.date = updated.date
            tx.description = updated.description
            tx.amount_cents = updated.amount_cents
            tx.category = updated.category
            tx.is_income = updated.is_income
            tx.member_id = updated.member_id
            tx.fingerprint = updated.fingerprint
            upsert_category(current_user_id(), tx.category)
            db.session.commit()
            flash("Transaction updated.", "success")
            return _safe_redirect("main.transactions")
        except ValidationError as exc:
            db.session.rollback()
            flash(str(exc), "error")
    categories = Category.query.filter_by(user_id=current_user_id()).order_by(Category.name.asc()).all()
    return render_template("transaction_edit.html", tx=tx, categories=categories, members=_member_rows(), page_title="Edit Transaction")


@bp.route("/transactions/<int:tx_id>/delete", methods=["POST"])
@login_required
@rate_limit(limit=60, window_seconds=60, label="transaction_delete")
def delete_transaction(tx_id: int):
    tx: Transaction = _get_owned(Transaction, tx_id)
    db.session.delete(tx)
    db.session.commit()
    flash("Transaction deleted.", "success")
    return _safe_redirect("main.transactions")


@bp.route("/transactions/import/preview", methods=["POST"])
@login_required
@rate_limit(limit=8, window_seconds=60, label="csv_preview")
def import_preview():
    uploaded = request.files.get("csv_file")
    if not uploaded or not uploaded.filename.lower().endswith(".csv"):
        flash("Upload a CSV file.", "error")
        return _safe_redirect("main.transactions")
    try:
        preview = preview_csv(uploaded)
        from flask import session

        session["csv_import_token"] = save_csv_preview_text(current_user_id(), preview["raw_text"])
        preview.pop("raw_text", None)
        return render_template("import_preview.html", preview=preview, page_title="CSV Import Preview")
    except ValidationError as exc:
        flash(str(exc), "error")
        return _safe_redirect("main.transactions")


@bp.route("/transactions/import/commit", methods=["POST"])
@login_required
@rate_limit(limit=8, window_seconds=60, label="csv_commit")
def import_commit():
    from flask import session

    token = session.pop("csv_import_token", "")
    try:
        text = load_csv_preview_text(current_user_id(), token)
    except ValidationError as exc:
        flash(str(exc), "error")
        return _safe_redirect("main.transactions")
    mapping = {
        "date": request.form.get("col_date", ""),
        "description": request.form.get("col_description", ""),
        "amount": request.form.get("col_amount", ""),
        "debit": request.form.get("col_debit", ""),
        "credit": request.form.get("col_credit", ""),
        "category": request.form.get("col_category", ""),
    }
    if not mapping["date"] or not mapping["description"] or not (mapping["amount"] or mapping["debit"] or mapping["credit"]):
        flash("Map date, description, and either amount or debit/credit columns.", "error")
        return _safe_redirect("main.transactions")
    result = import_csv_text(current_user_id(), g.workspace.id, text, mapping)
    flash(f"CSV imported: {result['imported']} added, {result['skipped']} skipped, {result['errors']} row error(s).", "success")
    return _safe_redirect("main.transactions")


@bp.route("/budgets", methods=["GET", "POST"])
@login_required
@rate_limit(limit=60, window_seconds=60, label="budgets")
def budgets():
    user_id = current_user_id()
    if request.method == "POST":
        try:
            category = normalize_category(request.form.get("category"))
            month = int(request.form.get("month") or "0")
            year = int(request.form.get("year") or "0")
            if not (1 <= month <= 12):
                raise ValidationError("Month must be 1-12.")
            if not (1900 <= year <= 2200):
                raise ValidationError("Year is out of range.")
            amount_cents = parse_money_to_cents(request.form.get("amount"), allow_negative=False, field_name="Budget amount")
            existing = Budget.query.filter_by(
                user_id=user_id, workspace_id=g.workspace.id, category=category, month=month, year=year
            ).first()
            if existing:
                existing.amount_cents = amount_cents
                flash("Budget updated.", "success")
            else:
                db.session.add(
                    Budget(
                        user_id=user_id, workspace_id=g.workspace.id, category=category, month=month, year=year, amount_cents=amount_cents
                    )
                )
                flash("Budget created.", "success")
            upsert_category(user_id, category)
            db.session.commit()
        except (ValidationError, ValueError) as exc:
            db.session.rollback()
            flash(str(exc), "error")
        return _safe_redirect("main.budgets")
    rows = budget_status_rows(user_id, g.workspace.id)
    categories = Category.query.filter_by(user_id=user_id).order_by(Category.name.asc()).all()
    today = date.today()
    return render_template("budgets.html", rows=rows, categories=categories, today=today, page_title="Budgets")


@bp.route("/budgets/<int:budget_id>/delete", methods=["POST"])
@login_required
@rate_limit(limit=50, window_seconds=60, label="budget_delete")
def delete_budget(budget_id: int):
    budget: Budget = _get_owned(Budget, budget_id)
    db.session.delete(budget)
    db.session.commit()
    flash("Budget deleted.", "success")
    return _safe_redirect("main.budgets")


@bp.route("/rules", methods=["GET", "POST"])
@login_required
@rate_limit(limit=60, window_seconds=60, label="rules")
def rules():
    user_id = current_user_id()
    if request.method == "POST":
        try:
            operator = request.form.get("operator", "contains")
            if operator not in {"contains", "equals", "starts_with", "ends_with"}:
                abort(400)
            value = clean_text(request.form.get("value"), max_length=80, field_name="Rule value")
            if not value:
                raise ValidationError("Rule value is required.")
            category = normalize_category(request.form.get("category"))
            min_raw = request.form.get("min_amount", "").strip()
            max_raw = request.form.get("max_amount", "").strip()
            min_cents = parse_money_to_cents(min_raw, allow_negative=False, field_name="Minimum amount") if min_raw else None
            max_cents = parse_money_to_cents(max_raw, allow_negative=False, field_name="Maximum amount") if max_raw else None
            if min_cents is not None and max_cents is not None and min_cents > max_cents:
                raise ValidationError("Minimum amount cannot exceed maximum amount.")
            db.session.add(
                Rule(
                    user_id=user_id,
                    field="description",
                    operator=operator,
                    value=value,
                    category=category,
                    min_cents=min_cents,
                    max_cents=max_cents,
                )
            )
            upsert_category(user_id, category)
            db.session.commit()
            flash("Rule created.", "success")
        except ValidationError as exc:
            db.session.rollback()
            flash(str(exc), "error")
        return _safe_redirect("main.rules")
    rows = Rule.query.filter_by(user_id=user_id).order_by(Rule.created_at.desc()).all()
    categories = Category.query.filter_by(user_id=user_id).order_by(Category.name.asc()).all()
    return render_template("rules.html", rows=rows, categories=categories, page_title="Rules")


@bp.route("/rules/<int:rule_id>/delete", methods=["POST"])
@login_required
@rate_limit(limit=50, window_seconds=60, label="rule_delete")
def delete_rule(rule_id: int):
    rule: Rule = _get_owned(Rule, rule_id)
    db.session.delete(rule)
    db.session.commit()
    flash("Rule deleted.", "success")
    return _safe_redirect("main.rules")


@bp.route("/categories", methods=["GET", "POST"])
@login_required
@rate_limit(limit=60, window_seconds=60, label="categories")
def categories():
    user_id = current_user_id()
    if request.method == "POST":
        try:
            upsert_category(user_id, request.form.get("name"))
            db.session.commit()
            flash("Category saved.", "success")
        except (ValidationError, IntegrityError) as exc:
            db.session.rollback()
            flash(str(exc), "error")
        return _safe_redirect("main.categories")
    rows = Category.query.filter_by(user_id=user_id).order_by(Category.name.asc()).all()
    return render_template("categories.html", rows=rows, page_title="Categories")


@bp.route("/categories/<int:category_id>/delete", methods=["POST"])
@login_required
@rate_limit(limit=50, window_seconds=60, label="category_delete")
def delete_category(category_id: int):
    category: Category = _get_owned(Category, category_id)
    if category.name in {"Income", "Uncategorized"}:
        flash("Core categories cannot be deleted.", "error")
        return _safe_redirect("main.categories")
    db.session.delete(category)
    db.session.commit()
    flash("Category deleted.", "success")
    return _safe_redirect("main.categories")


@bp.route("/analyst")
@login_required
def analyst():
    run = AnalystRun.query.filter_by(user_id=current_user_id(), workspace_id=g.workspace.id).order_by(AnalystRun.created_at.desc()).first()
    result = parse_analyst_result(run)
    return render_template("analyst.html", run=run, result=result, page_title="Analyst")


@bp.route("/analyst/run", methods=["POST"])
@login_required
@rate_limit(limit=8, window_seconds=60, label="analyst_run")
def analyst_run():
    generate_analyst_report(current_user_id(), g.workspace.id, settings=g.settings)
    flash("Full analyst run completed.", "success")
    return _safe_redirect("main.analyst")


@bp.route("/ai")
@login_required
def ai_page():
    summary = dashboard_summary(current_user_id(), g.workspace.id, settings=g.settings, workspace=g.workspace)
    return render_template("ai.html", summary=summary, page_title="Ouroboros Advisor")


@bp.route("/api/advisor/chat", methods=["POST"])
@login_required
@rate_limit(limit=50, window_seconds=60, label="advisor_chat")
def advisor_chat():
    payload = request.get_json(silent=True) or {}
    try:
        response = ouroboros_advisor_reply(current_user_id(), g.workspace.id, str(payload.get("message") or ""), settings=g.settings)
        return jsonify({"ok": True, **response})
    except ValidationError as exc:
        return jsonify({"ok": False, "error": str(exc)}), 400


@bp.route("/toolkit")
@login_required
def toolkit():
    return render_template("toolkit.html", page_title="Planning Toolkit")


@bp.route("/api/currency/rate")
@login_required
@rate_limit(limit=40, window_seconds=60, label="fx")
def currency_rate():
    base = str(request.args.get("base") or g.settings.base_currency_code or "USD").upper()
    target = str(request.args.get("target") or g.settings.display_currency_code or base).upper()
    try:
        data = get_exchange_rate(base, target)
        data["symbol"] = CURRENCY_OPTIONS.get(target, "")
        return jsonify({"ok": True, **data})
    except ValidationError as exc:
        return jsonify({"ok": False, "error": str(exc), "base": base, "target": target, "symbol": CURRENCY_OPTIONS.get(target, "")}), 503


@bp.route("/calendar")
@login_required
def calendar_view():
    today = date.today()
    try:
        year = int(request.args.get("year", today.year))
        month = int(request.args.get("month", today.month))
        if not (1900 <= year <= 2200 and 1 <= month <= 12):
            raise ValueError
    except ValueError:
        year, month = today.year, today.month
    prev_y, prev_m = previous_month(year, month)
    next_y, next_m = next_month(year, month)
    data = calendar_data(current_user_id(), g.workspace.id, year=year, month=month, settings=g.settings)
    return render_template(
        "calendar.html", data=data, year=year, month=month, prev=(prev_y, prev_m), next=(next_y, next_m), page_title="Calendar"
    )


@bp.route("/reports")
@login_required
def reports():
    user_id = current_user_id()
    transactions = Transaction.query.filter_by(user_id=user_id, workspace_id=g.workspace.id).order_by(Transaction.date.desc()).all()
    total_income = sum(t.amount_cents for t in transactions if t.is_income)
    total_expense = sum(abs(t.amount_cents) for t in transactions if not t.is_income)
    net = total_income - total_expense
    category_totals: dict[str, int] = {}
    for t in transactions:
        if not t.is_income:
            category_totals[t.category] = category_totals.get(t.category, 0) + abs(t.amount_cents)
    budget_rows = budget_status_rows(user_id, g.workspace.id)
    return render_template(
        "reports.html",
        rows=transactions,
        rows_json=_transaction_json(transactions[:5000]),
        total_income=total_income,
        total_expense=total_expense,
        net=net,
        category_totals=category_totals,
        budget_rows=budget_rows,
        page_title="Reports",
    )


def _report_data(user_id: int):
    transactions = Transaction.query.filter_by(user_id=user_id, workspace_id=g.workspace.id).order_by(Transaction.date.desc()).all()
    total_income = sum(t.amount_cents for t in transactions if t.is_income)
    total_expense = sum(abs(t.amount_cents) for t in transactions if not t.is_income)
    return transactions, total_income, total_expense, total_income - total_expense


@bp.route("/download/<fmt>")
@login_required
@rate_limit(limit=20, window_seconds=60, label="download")
def download_report(fmt: str):
    rows, total_income, total_expense, net = _report_data(current_user_id())
    if fmt == "xlsx":
        wb = Workbook()
        ws = wb.active
        ws.title = "Transactions"
        ws.append(["Workspace", "Date", "Description", "Member", "Category", "Type", "Amount"])
        for t in rows:
            ws.append(
                [
                    g.workspace.name,
                    t.date.isoformat(),
                    spreadsheet_safe(t.description),
                    spreadsheet_safe(t.member.name if t.member else ""),
                    spreadsheet_safe(t.category),
                    "Income" if t.is_income else "Expense",
                    Decimal(t.amount_cents) / Decimal(100),
                ]
            )
        summary = wb.create_sheet("Summary")
        summary.append(["Metric", "Value"])
        summary.append(["Total Income", Decimal(total_income) / Decimal(100)])
        summary.append(["Total Expenses", Decimal(total_expense) / Decimal(100)])
        summary.append(["Net Cash Flow", Decimal(net) / Decimal(100)])
        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return send_file(
            buffer,
            as_attachment=True,
            download_name="financial_report.xlsx",
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
    if fmt == "docx":
        from .money import format_cents

        symbol = CURRENCY_OPTIONS.get(g.settings.base_currency_code, "$")
        code = g.settings.base_currency_code
        doc = Document()
        doc.add_heading("Financial Report", 0)
        doc.add_paragraph(f"Workspace: {g.workspace.name} ({g.workspace.workspace_type})")
        doc.add_paragraph(f"Total Income: {format_cents(total_income, symbol, code)}")
        doc.add_paragraph(f"Total Expenses: {format_cents(total_expense, symbol, code)}")
        doc.add_paragraph(f"Net Cash Flow: {format_cents(net, symbol, code)}")
        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        headers = ["Date", "Description", "Member", "Category", "Type", "Amount"]
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        for t in rows[:200]:
            cells = table.add_row().cells
            cells[0].text = t.date.isoformat()
            cells[1].text = t.description
            cells[2].text = t.member.name if t.member else ""
            cells[3].text = t.category
            cells[4].text = "Income" if t.is_income else "Expense"
            cells[5].text = format_cents(t.amount_cents, symbol, code)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return send_file(
            buffer,
            as_attachment=True,
            download_name="financial_report.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    if fmt == "pdf":
        try:
            from weasyprint import HTML
        except Exception:
            abort(500, description="PDF export requires WeasyPrint system dependencies.")
        html = render_template("pdf_report.html", rows=rows[:100], total_income=total_income, total_expense=total_expense, net=net)
        pdf = HTML(string=html).write_pdf()
        return send_file(io.BytesIO(pdf), as_attachment=True, download_name="financial_report.pdf", mimetype="application/pdf")
    abort(404)


@bp.route("/help")
@login_required
def help_page():
    return render_template("help.html", page_title="Help")


@bp.app_errorhandler(400)
def bad_request(error):
    return render_template("error.html", code=400, message="Bad request or invalid security token."), 400


@bp.app_errorhandler(403)
def forbidden(error):
    return render_template("error.html", code=403, message="Blocked by localhost-only security."), 403


@bp.app_errorhandler(404)
def not_found(error):
    return render_template("error.html", code=404, message="Page or record not found."), 404


@bp.app_errorhandler(429)
def too_many(error):
    return render_template("error.html", code=429, message="Rate limit reached. Wait a moment and retry."), 429


@bp.app_errorhandler(500)
def server_error(error):
    return render_template("error.html", code=500, message=getattr(error, "description", "Local server error.")), 500
